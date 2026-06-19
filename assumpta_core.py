from __future__ import annotations

import html
import io
import json
import os
import re
import subprocess
import textwrap
import urllib.parse
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from urllib.error import URLError
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


BASE_DIR = os.path.dirname(__file__)
STATE_PATH = os.path.join(BASE_DIR, "state.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
PREVIEW_DIR = os.path.join(BASE_DIR, "static", "previews")
ASSET_DIR = os.path.join(BASE_DIR, "assets")
LITURGY_SOURCE_DIR = os.path.join(ASSET_DIR, "liturgy_sources")
TEMPLATE_PAGE_1 = os.path.join(ASSET_DIR, "template", "assumpta-template-1.png")
TEMPLATE_PAGE_2 = os.path.join(ASSET_DIR, "template", "assumpta-template-2.png")
PDFTOPPM = "/Users/javiermanuelrodriguezrodriguez/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/pdftoppm"


DEFAULT_STATE = {
    "last_assumpta_number": 1233,
    "last_assumpta_date": "2026-06-07",
    "next_compendio_question": 330,
}

COMPENDIO_HEADER = """SEGUNDA SECCION
LOS SIETE SACRAMENTOS DE LA IGLESIA
CAPITULO SEGUNDO
LOS SACRAMENTOS AL SERVICIO DE LA COMUNION Y DE LA MISION

EL SACRAMENTO DEL ORDEN"""


SAINTS_BY_MONTH_DAY = {
    "06-11": ["Bernabe"],
    "06-13": ["Antonio de Padua"],
    "06-14": ["San Eliseo, profeta", "Santos Anastasio, Felix y Digna, martires"],
    "06-15": ["Santa Micaela del Santisimo Sacramento", "San Bernardo de Menthon"],
    "06-16": ["San Aureliano, obispo"],
    "06-17": ["San Gregorio Barbarigo", "San Ismael, martir"],
    "06-18": ["Santa Isabel de Schönau", "San Ciriaco, martir"],
    "06-19": ["San Romualdo, abad", "San Gervasio, martir"],
    "06-20": ["Santa Florentina", "San Silverio, papa y martir"],
}


@dataclass
class GospelData:
    date_label: str
    celebration: str
    reading_ref: str
    gospel_ref: str
    gospel_text: str
    source_url: str
    aleluya: str = ""
    warning: str = ""


SPANISH_MONTHS = [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]


def load_state() -> dict:
    if not os.path.exists(STATE_PATH):
        save_state(DEFAULT_STATE)
    with open(STATE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_state(data: dict) -> None:
    with open(STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def next_assumpta_number() -> int:
    return int(load_state().get("last_assumpta_number", 1233)) + 1


def parse_date(value: str) -> date:
    return datetime.strptime(value, "%Y-%m-%d").date()


def long_spanish_date(value: str) -> str:
    months = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    d = parse_date(value)
    return f"{d.day} de {months[d.month - 1]} de {d.year}"


def saints_for_week(start_iso: str) -> list[dict]:
    start = parse_date(start_iso)
    result = []
    for offset in range(7):
        day = start + timedelta(days=offset)
        key = day.strftime("%m-%d")
        for saint in SAINTS_BY_MONTH_DAY.get(key, []):
            result.append({"date": day.isoformat(), "name": saint})
    return result


def extract_docx_text(file_bytes: bytes) -> str:
    if not file_bytes:
        return ""
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_compendio_questions(pdf_path: str, start_number: int, max_chars: int = 1800) -> tuple[str, int]:
    if not pdf_path or not os.path.exists(pdf_path):
        return f"No se encontro el Compendio en: {pdf_path}", start_number
    reader = PdfReader(pdf_path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\d{1,2}/\d{1,2}/\d{2},\s*\d{1,2}:\d{2}\s*about:blank", " ", text)
    text = re.sub(r"about:blank\s+\d+/\d+", " ", text)
    questions = []
    current = start_number
    while current < 600:
        pattern = rf"(?ms)(^|\n){current}\.\s+(.+?)(?=\n{current + 1}\.\s+|\nCAP[IÍ]TULO|\n[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ ]{{8,}}|\Z)"
        match = re.search(pattern, text)
        if not match:
            break
        block = re.sub(r"\n{3,}", "\n\n", f"{current}. {match.group(2).strip()}").strip()
        block = block.replace("ﬁ", "fi").replace("ﬂ", "fl")
        block = re.sub(r"(\?)\s+(?:\d{3,4}(?:-\d{3,4})?\s*)+", r"\1\n", block)
        candidate = "\n\n".join(questions + [block])
        if len(candidate) > max_chars and questions:
            break
        questions.append(block)
        current += 1
    if not questions:
        return f"No se encontro la pregunta {start_number} en el Compendio.", start_number
    return COMPENDIO_HEADER + "\n\n" + "\n\n".join(questions), current


def generate_pastor_letter(celebration: str, gospel_ref: str) -> str:
    theme = celebration if celebration and "pendiente" not in celebration.lower() else "la liturgia de este domingo"
    reference = f" ({gospel_ref})" if gospel_ref and "pendiente" not in gospel_ref.lower() else ""
    return (
        f"Celebramos {theme}. La Palabra de Dios{reference} nos invita a renovar la fe "
        "con sencillez y a poner a Cristo en el centro de nuestra vida cotidiana. "
        "Pidamos al Senor un corazon disponible para escucharle, recibirle en la Eucaristia "
        "y servir con alegria a quienes tenemos cerca. Que esta semana sea una ocasion "
        "para crecer como comunidad parroquial, unidos en la oracion, la caridad y la esperanza."
    )


def gospel_title_from_ref(gospel_ref: str) -> str:
    ref = (gospel_ref or "").strip()
    books = {
        "Mateo": "san Mateo",
        "Marcos": "san Marcos",
        "Lucas": "san Lucas",
        "Juan": "san Juan",
    }
    for book, liturgical in books.items():
        if ref.startswith(book):
            return "Santo Evangelio según " + ref.replace(book, liturgical, 1)
    if ref and "pendiente" not in ref.lower():
        return "Santo Evangelio según la referencia indicada: " + ref
    return "Santo Evangelio según la referencia indicada"


def _archimadrid_month_urls(d: date) -> list[str]:
    month_name = SPANISH_MONTHS[d.month - 1]
    upload_months = [(d.month - 1) or 12, d.month]
    urls = []
    for upload_month in upload_months:
        upload_year = d.year - 1 if d.month == 1 and upload_month == 12 else d.year
        urls.append(
            "https://oracionyliturgia.archimadrid.org/wp-content/uploads/"
            f"{upload_year}/{upload_month:02d}/{month_name}_{d.year}.pdf"
        )
    return list(dict.fromkeys(urls))


def _download_archimadrid_pdf(d: date) -> tuple[str, str]:
    os.makedirs(LITURGY_SOURCE_DIR, exist_ok=True)
    month_name = SPANISH_MONTHS[d.month - 1]
    cache_path = os.path.join(LITURGY_SOURCE_DIR, f"{month_name}_{d.year}.pdf")
    if os.path.exists(cache_path):
        return cache_path, _archimadrid_month_urls(d)[0]
    last_error = ""
    for url in _archimadrid_month_urls(d):
        try:
            req = Request(url, headers={"User-Agent": "AssumptaGenerator/1.0"})
            with urlopen(req, timeout=25) as response:
                data = response.read()
            if data.startswith(b"%PDF"):
                with open(cache_path, "wb") as f:
                    f.write(data)
                return cache_path, url
        except Exception as exc:
            last_error = str(exc)
    raise URLError(last_error or "No se pudo descargar el PDF mensual de Archimadrid.")


def _clean_liturgy_text(text: str) -> str:
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"www\.archimadrid\.org", " ", text)
    text = re.sub(r"Dpto\. Internet Arzobispado de Madrid", " ", text)
    text = re.sub(r"Lecturas de la Misa", " ", text)
    text = re.sub(r"webmaster@archimadrid\.org", " ", text)
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if re.match(r"^[◆\-\s]*\d{1,3}\s*-\s*\d{1,3}[◆\-\s]*$", stripped):
            continue
        if stripped in SPANISH_MONTHS or re.match(r"^[◆\s]+$", stripped):
            continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_reference_from_title(title: str) -> str:
    match = re.search(r"según san\s+([A-Za-zÁÉÍÓÚáéíóúÑñ]+)\s+(.+?)(?:\.|$)", title)
    if not match:
        return "Evangelio pendiente"
    book = match.group(1).strip()
    ref = match.group(2).strip()
    return f"{book} {ref}".replace("  ", " ")


def _extract_archimadrid_block(text: str, d: date) -> str:
    date_pattern = rf"{d.day}/{d.month}/{d.year}\s+-"
    start_match = re.search(date_pattern, text)
    if not start_match:
        return ""
    next_day = d + timedelta(days=1)
    next_pattern = rf"\n{next_day.day}/{next_day.month}/{next_day.year}\s+-"
    next_match = re.search(next_pattern, text[start_match.end():])
    end = start_match.end() + next_match.start() if next_match else len(text)
    return _clean_liturgy_text(text[start_match.start():end])


def fetch_archimadrid_liturgy(fecha_iso: str) -> GospelData:
    d = parse_date(fecha_iso)
    try:
        pdf_path, url = _download_archimadrid_pdf(d)
        reader = PdfReader(pdf_path)
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        block = _extract_archimadrid_block(full_text, d)
        if not block:
            raise ValueError("No se encontro la fecha en el PDF mensual.")
    except Exception as exc:
        fallback_url = _archimadrid_month_urls(d)[0]
        return GospelData(
            date_label=fecha_iso,
            celebration="Celebracion pendiente",
            reading_ref="Lecturas pendientes",
            gospel_ref="Evangelio pendiente",
            gospel_text="No se pudo extraer la liturgia de Archimadrid. Pegue el Evangelio manualmente.",
            source_url=fallback_url,
            aleluya="",
            warning=f"Fuente Archimadrid no disponible: {exc}",
        )

    lines = [line.strip() for line in block.splitlines() if line.strip()]
    celebration = ""
    if lines:
        first = lines[0]
        celebration = re.sub(r"^\d{1,2}/\d{1,2}/\d{4}\s+-\s*", "", first).strip()
        extras = []
        for line in lines[1:4]:
            if re.match(r"^(1ª lectura|Salmo|2ª lectura|SECUENCIA|Aleluya|Evangelio):?", line):
                break
            extras.append(line)
        if extras:
            celebration = " ".join([celebration] + extras)

    reading_parts = []
    for prefix in ("1ª lectura:", "Salmo:", "2ª lectura:"):
        for line in lines:
            if line.startswith(prefix):
                reading_parts.append(line)
                break

    aleluya = ""
    gospel_heading = ""
    gospel_title = ""
    gospel_text = ""
    try:
        ev_index = next(i for i, line in enumerate(lines) if line.startswith("Evangelio:"))
        gospel_heading = lines[ev_index]
        title_index = next(i for i in range(ev_index + 1, len(lines)) if "Evangelio según san" in lines[i])
        gospel_title = lines[title_index].rstrip(".")
        gospel_text = "\n".join(lines[title_index + 1:]).strip()
        aleluya_start = max((i for i, line in enumerate(lines[:ev_index]) if line.startswith("Aleluya")), default=-1)
        if aleluya_start >= 0:
            aleluya = "\n".join(lines[aleluya_start:ev_index]).strip()
    except (StopIteration, ValueError):
        gospel_text = "No se pudo extraer automaticamente el Evangelio."

    gospel_ref = _extract_reference_from_title(gospel_title)
    return GospelData(
        date_label=fecha_iso,
        celebration=celebration or "Celebracion pendiente",
        reading_ref="     ".join(reading_parts) or "Lecturas pendientes",
        gospel_ref=gospel_ref,
        gospel_text=gospel_text,
        source_url=url,
        aleluya=aleluya,
        warning="Fuente: Oracion y Liturgia - Archimadrid.",
    )


def fetch_vatican_gospel(fecha_iso: str) -> GospelData:
    return fetch_archimadrid_liturgy(fecha_iso)


def suggested_liturgy_image_query(gospel: GospelData) -> str:
    text = " ".join([gospel.celebration, gospel.gospel_ref, gospel.gospel_text[:300]]).lower()
    if "pan vivo" in text or "carne" in text or "sangre" in text or "cuerpo de cristo" in text:
        return "Eucharist chalice bread"
    if "buen pastor" in text or "ovejas" in text:
        return "Good Shepherd Jesus"
    if "bienaventurados" in text or "monte" in text:
        return "Sermon on the Mount"
    if "padre nuestro" in text:
        return "Jesus teaching prayer"
    if "luz del mundo" in text:
        return "Christ light of the world"
    if "cruz" in text:
        return "Jesus carrying the cross"
    return gospel.celebration or gospel.gospel_ref or "Gospel"


def _safe_name(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip().lower()).strip("-")[:60] or "imagen"


def download_commons_image(query: str, folder: str = "downloaded") -> str:
    os.makedirs(os.path.join(ASSET_DIR, folder), exist_ok=True)
    filename = os.path.join(ASSET_DIR, folder, _safe_name(query) + ".jpg")
    if os.path.exists(filename):
        return filename

    api = "https://commons.wikimedia.org/w/api.php?" + urllib.parse.urlencode({
        "action": "query",
        "format": "json",
        "generator": "search",
        "gsrnamespace": "6",
        "gsrsearch": query,
        "gsrlimit": "1",
        "prop": "imageinfo",
        "iiprop": "url",
    })
    try:
        with urlopen(Request(api, headers={"User-Agent": "AssumptaGenerator/1.0"}), timeout=20) as response:
            payload = json.loads(response.read().decode("utf-8"))
        pages = payload.get("query", {}).get("pages", {})
        for page in pages.values():
            infos = page.get("imageinfo") or []
            if infos and infos[0].get("url"):
                image_url = infos[0]["url"]
                with urlopen(Request(image_url, headers={"User-Agent": "AssumptaGenerator/1.0"}), timeout=20) as image_response:
                    data = image_response.read()
                with open(filename, "wb") as f:
                    f.write(data)
                return filename
    except Exception:
        return ""
    return ""


def _fit_font_size(text: str, base_size: int, min_size: int, max_chars: int) -> int:
    length = len(text or "")
    if length <= max_chars:
        return base_size
    overflow = min(1.0, (length - max_chars) / max_chars)
    return max(min_size, int(base_size - overflow * 5))


def _draw_cover(c: canvas.Canvas, x: float, y: float, w: float, h: float, fill=colors.white) -> None:
    c.setFillColor(fill)
    c.setStrokeColor(fill)
    c.rect(x, y, w, h, stroke=0, fill=1)
    c.setFillColor(colors.black)
    c.setStrokeColor(colors.black)


def _wrap_lines(text: str, width: float, font_name: str, font_size: int) -> list[str]:
    avg = max(10, int(width / max(stringWidth("n", font_name, font_size), 1)))
    lines = []
    for paragraph in (text or "").splitlines():
        paragraph = paragraph.strip()
        if not paragraph:
            lines.append("")
            continue
        lines.extend(textwrap.wrap(paragraph, width=avg))
    return lines


def _draw_wrapped(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    width: float,
    height: float,
    font_size: int,
    font_name: str = "Helvetica",
    align: str = "left",
    leading_factor: float = 1.16,
    italic: bool = False,
) -> bool:
    if italic and font_name == "Helvetica":
        font_name = "Helvetica-Oblique"
    c.setFont(font_name, font_size)
    line_height = font_size * leading_factor
    cursor = y + height - font_size
    ok = True
    for line in _wrap_lines(text, width, font_name, font_size):
        if cursor < y:
            ok = False
            break
        if not line:
            cursor -= line_height * 0.7
            continue
        if align == "center":
            c.drawCentredString(x + width / 2, cursor, line)
        elif align == "right":
            c.drawRightString(x + width, cursor, line)
        else:
            c.drawString(x, cursor, line)
        cursor -= line_height
    return ok


def _draw_overflow_warning(c: canvas.Canvas, x: float, y: float, width: float, message: str) -> None:
    c.setFillColor(colors.HexColor("#b00020"))
    c.setFont("Helvetica-Bold", 6)
    c.drawString(x, y, message)
    c.setFillColor(colors.black)


def _draw_image_or_placeholder(c: canvas.Canvas, path: str, x: float, y: float, w: float, h: float, label: str) -> None:
    if path and os.path.exists(path):
        try:
            c.drawImage(path, x, y, w, h, preserveAspectRatio=True, anchor="c")
            return
        except Exception:
            pass
    _draw_cover(c, x, y, w, h, colors.HexColor("#eeeeee"))
    c.setStrokeColor(colors.HexColor("#bdbdbd"))
    c.rect(x, y, w, h, stroke=1, fill=0)
    c.setFillColor(colors.HexColor("#555555"))
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(x + w / 2, y + h / 2, label[:26])
    c.setFillColor(colors.black)


def _draw_announcement(c: canvas.Canvas, data: dict, warnings: list[str]) -> None:
    _draw_cover(c, 331, 114, 184, 270)
    text = data.get("anuncio", "Anuncio pendiente.").strip()
    image_path = data.get("anuncio_imagen_path", "").strip()
    has_image = bool(image_path and os.path.exists(image_path))
    content_x, content_y, content_w, content_h = 342, 132, 162, 226

    if has_image and text:
        image_h = 122
        text_h = 86
        _draw_image_or_placeholder(c, image_path, content_x, content_y + text_h + 14, content_w, image_h, "Imagen del anuncio")
        text_size = _fit_font_size(text, 11, 7, 360)
        ok = _draw_wrapped(c, text, content_x, content_y, content_w, text_h, text_size, "Helvetica-Bold", align="center", leading_factor=1.16)
    elif has_image:
        _draw_image_or_placeholder(c, image_path, content_x, content_y + 6, content_w, content_h - 12, "Imagen del anuncio")
        ok = True
    else:
        text_size = _fit_font_size(text, 16, 8, 600)
        ok = _draw_wrapped(c, text, content_x, 150, content_w, 190, text_size, "Helvetica-Bold", align="center", leading_factor=1.18)

    if not ok:
        warnings.append("Anuncio central: el contenido es demasiado largo para esta zona.")
        _draw_overflow_warning(c, 339, 132, 165, "Anuncio demasiado largo")


def _selected_saint_names(data: dict) -> list[str]:
    value = data.get("santos", "")
    if isinstance(value, list):
        return value[:2]
    return [line.strip() for line in str(value).splitlines() if line.strip()][:2]


def _draw_page_1(c: canvas.Canvas, data: dict, page_w: float, page_h: float, warnings: list[str]) -> None:
    c.drawImage(TEMPLATE_PAGE_1, 0, 0, page_w, page_h)

    # Catecismo: panel izquierdo, respetando la cabecera gris original.
    _draw_cover(c, 19, 24, 241, 472)
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(140, 474, "CATECISMO DE LA IGLESIA CATOLICA")
    size = _fit_font_size(data["compendio"], 7, 5, 2650)
    ok = _draw_wrapped(c, data["compendio"], 21, 33, 235, 430, size, "Helvetica", leading_factor=1.12)
    if not ok:
        warnings.append("Catecismo: el texto es demasiado largo para esta zona.")
        _draw_overflow_warning(c, 22, 27, 235, "Catecismo demasiado largo")

    # Portada: solo numero, fecha y carta. El titulo ASSUMPTA se conserva del fondo.
    _draw_cover(c, 656, 446, 172, 74)
    _draw_cover(c, 656, 500, 172, 25)
    c.setFillColor(colors.HexColor("#555555"))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(659, 486, f"NUMERO {data['numero']}")
    c.drawString(659, 468, data["fecha_larga"])
    c.setFillColor(colors.black)

    _draw_cover(c, 656, 154, 168, 318)
    c.setFont("Helvetica-Oblique", 10)
    c.drawString(659, 446, "Queridos feligreses:")
    carta_size = _fit_font_size(data["carta"], 10, 7, 1350)
    ok = _draw_wrapped(c, data["carta"], 659, 174, 160, 265, carta_size, "Helvetica", align="left", italic=True, leading_factor=1.28)
    c.setFont("Helvetica-Oblique", 10)
    c.drawRightString(819, 164, "Vuestro Parroco")
    if not ok:
        warnings.append("Carta del parroco: el texto es demasiado largo para esta zona.")
        _draw_overflow_warning(c, 659, 156, 160, "Carta demasiado larga")

    # Santos confirmados: imagenes automaticas y nombres.
    _draw_cover(c, 656, 26, 168, 126)
    names = _selected_saint_names(data)
    slots = [(660, 72, 78, 72), (742, 72, 78, 72)]
    for index, (name, slot) in enumerate(zip(names, slots)):
        image_path = download_commons_image(name, "saints")
        _draw_image_or_placeholder(c, image_path, *slot, label=name)
        c.setFont("Helvetica", 6.5)
        c.drawCentredString(slot[0] + slot[2] / 2, 60, name[:24])
    if not names:
        c.setFont("Helvetica", 8)
        c.drawCentredString(740, 105, "Santos pendientes de elegir")


def _draw_page_2(c: canvas.Canvas, data: dict, page_w: float, page_h: float, warnings: list[str]) -> None:
    c.drawImage(TEMPLATE_PAGE_2, 0, 0, page_w, page_h)

    # Liturgia completa: imagen, titulo, lecturas, secuencia/aleluya y Evangelio.
    _draw_cover(c, 18, 23, 259, 551)
    lit_image = data.get("liturgia_imagen_path") or download_commons_image(data.get("liturgia_imagen_query") or data.get("celebracion") or data.get("evangelio_ref") or "Gospel", "liturgy")
    _draw_image_or_placeholder(c, lit_image, 18, 446, 259, 110, "Imagen del Evangelio")
    c.setFillColor(colors.HexColor("#666666"))
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(147, 431, data.get("celebracion", "Domingo").replace(" - ", "\n")[:42])
    c.setFont("Helvetica-BoldOblique", 9)
    c.drawCentredString(147, 409, data.get("lecturas", "Lecturas pendientes"))
    c.setFillColor(colors.black)

    secuencia = data.get("secuencia", "").strip()
    aleluya = data.get("aleluya", "").strip()
    evangelio_titulo = data.get("evangelio_titulo", "").strip()
    lit_pre = "\n\n".join(part for part in [secuencia, aleluya, evangelio_titulo] if part)
    ok_pre = _draw_wrapped(c, lit_pre, 23, 268, 248, 109, 7, "Helvetica", align="center", leading_factor=1.12)
    ev_size = _fit_font_size(data["evangelio"], 8, 6, 2500)
    ok_ev = _draw_wrapped(c, data["evangelio"], 23, 50, 248, 210, ev_size, "Helvetica", leading_factor=1.13)
    if not (ok_pre and ok_ev):
        warnings.append("Evangelio/liturgia: el texto es demasiado largo para esta zona.")
        _draw_overflow_warning(c, 24, 34, 245, "Liturgia demasiado larga")

    # Hueco del anuncio central, dejando cabecera, avisos y pie existentes.
    _draw_announcement(c, data, warnings)

    # Habla el Papa: conservar marco y cabecera; tapar solo el texto central.
    _draw_cover(c, 592, 183, 218, 204)
    papa = data.get("habla_papa", "Texto pendiente de cargar desde Word.")
    papa_size = _fit_font_size(papa, 9, 6, 1150)
    ok = _draw_wrapped(c, papa, 596, 195, 210, 184, papa_size, "Helvetica", italic=True, leading_factor=1.15)
    if not ok:
        warnings.append("Habla el Papa: el texto es demasiado largo para esta zona.")
        _draw_overflow_warning(c, 596, 187, 210, "Habla el Papa demasiado largo")


def _draw_proof_mark(c: canvas.Canvas, page_w: float, page_h: float) -> None:
    c.saveState()
    c.setFillColor(colors.Color(0.8, 0.1, 0.1, alpha=0.18))
    c.setFont("Helvetica-Bold", 72)
    c.translate(page_w / 2, page_h / 2)
    c.rotate(25)
    c.drawCentredString(0, 0, "PRUEBA")
    c.restoreState()


def build_pdf(data: dict, output_path: str, proof: bool = False) -> list[str]:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    c = canvas.Canvas(output_path, pagesize=landscape(A4))
    page_w, page_h = landscape(A4)
    warnings: list[str] = []
    _draw_page_1(c, data, page_w, page_h, warnings)
    if proof:
        _draw_proof_mark(c, page_w, page_h)
    c.showPage()
    _draw_page_2(c, data, page_w, page_h, warnings)
    if proof:
        _draw_proof_mark(c, page_w, page_h)
    c.save()
    return warnings


def build_previews(pdf_path: str) -> list[str]:
    os.makedirs(PREVIEW_DIR, exist_ok=True)
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    prefix = os.path.join(PREVIEW_DIR, base)
    if os.path.exists(PDFTOPPM):
        subprocess.run(
            [PDFTOPPM, "-jpeg", "-r", "110", pdf_path, prefix],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
    previews = []
    for index in (1, 2):
        path = f"{prefix}-{index}.jpg"
        if os.path.exists(path):
            previews.append("/static/previews/" + os.path.basename(path))
    return previews


def generate_assumpta(data: dict, proof: bool = False) -> dict:
    suffix = "prueba" if proof else "produccion"
    base = f"assumpta-{data['numero']}-{data['fecha']}-{suffix}.pdf"
    output = os.path.join(OUTPUT_DIR, base)
    warnings = build_pdf(data, output, proof=proof)
    previews = build_previews(output)
    return {"pdf": output, "previews": previews, "warnings": warnings}
